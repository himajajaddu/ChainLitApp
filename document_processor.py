import PyPDF2
import docx
import asyncio
import os
import tempfile
from typing import Optional, Any
import chainlit as cl

class DocumentProcessor:
    """Handles document processing for various file formats"""
    
    SUPPORTED_EXTENSIONS = {'.pdf', '.docx', '.txt'}
    
    def is_supported_file(self, filename: str) -> bool:
        """Check if the file type is supported"""
        return any(filename.lower().endswith(ext) for ext in self.SUPPORTED_EXTENSIONS)
    
    async def process_document(self, file_element: Any) -> Optional[str]:
        """Process uploaded document and extract text content"""
        try:
            # Get file extension
            file_extension = os.path.splitext(file_element.name)[1].lower()
            
            # Create temporary file
            with tempfile.NamedTemporaryFile(delete=False, suffix=file_extension, mode='wb') as temp_file:
                if hasattr(file_element, 'content') and file_element.content:
                    content = file_element.content
                    if isinstance(content, str):
                        content = content.encode('utf-8')
                    temp_file.write(content)
                elif hasattr(file_element, 'path'):
                    with open(file_element.path, 'rb') as src:
                        temp_file.write(src.read())
                temp_file_path = temp_file.name
            
            try:
                # Process based on file type
                if file_extension == '.pdf':
                    content = await self._process_pdf(temp_file_path)
                elif file_extension == '.docx':
                    content = await self._process_docx(temp_file_path)
                elif file_extension == '.txt':
                    content = await self._process_txt(temp_file_path)
                else:
                    raise ValueError(f"Unsupported file type: {file_extension}")
                
                return content
            
            finally:
                # Clean up temporary file
                try:
                    os.unlink(temp_file_path)
                except:
                    pass
        
        except Exception as e:
            raise Exception(f"Error processing document: {str(e)}")
    
    async def _process_pdf(self, file_path: str) -> str:
        """Extract text from PDF file"""
        try:
            content = []
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    page_text = page.extract_text()
                    if page_text.strip():
                        content.append(f"--- Page {page_num + 1} ---\n{page_text}")
            
            if not content:
                raise ValueError("No readable text found in PDF")
            
            return '\n\n'.join(content)
        
        except Exception as e:
            raise Exception(f"Error reading PDF: {str(e)}")
    
    async def _process_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            doc = docx.Document(file_path)
            content = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content.append(paragraph.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        content.append(' | '.join(row_text))
            
            if not content:
                raise ValueError("No readable text found in DOCX")
            
            return '\n'.join(content)
        
        except Exception as e:
            raise Exception(f"Error reading DOCX: {str(e)}")
    
    async def _process_txt(self, file_path: str) -> str:
        """Extract text from TXT file"""
        try:
            encodings = ['utf-8', 'utf-16', 'latin1', 'cp1252']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        content = file.read()
                        if content.strip():
                            return content
                except UnicodeDecodeError:
                    continue
            
            raise ValueError("Unable to decode text file with supported encodings")
        
        except Exception as e:
            raise Exception(f"Error reading TXT: {str(e)}")
    
    def get_document_stats(self, content: str) -> dict:
        """Get basic statistics about the document"""
        return {
            'character_count': len(content),
            'word_count': len(content.split()),
            'line_count': len(content.split('\n')),
            'paragraph_count': len([p for p in content.split('\n\n') if p.strip()])
        }
